from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "entregables"
OUT.mkdir(exist_ok=True)


BLUE = "1A365D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "DADCE0"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color=BORDER, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, name="Calibri", size=11, color=None, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_hyperlink_like(paragraph, label, url):
    run = paragraph.add_run(label + ": ")
    set_font(run, bold=True, color=BLUE)
    link = paragraph.add_run(url)
    set_font(link, color="0563C1")
    link.underline = True


def add_body_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    run = p.add_run(text)
    set_font(run)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run)
    return p


def add_numbered(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_font(run)
    return p


def style_table(table, header_fill=LIGHT_GRAY):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_font(run, size=10, bold=row_idx == 0)


def add_table(doc, headers, rows, widths, header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        hdr[i].width = Inches(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            cells[i].width = Inches(widths[i])
    style_table(table, header_fill)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    p.text = ""
    run = p.add_run(text)
    set_font(run, size=16 if level == 1 else 13, color=BLUE, bold=True)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 6)
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.add_run("TAREA - CONTROL DE VERSIONES")
set_font(title_run, size=20, color=BLUE, bold=True)
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle_run = subtitle.add_run("Practica guiada: Portafolio Web Express")
set_font(subtitle_run, size=12, color="555555")

meta = doc.add_table(rows=4, cols=2)
meta_rows = [
    ("Estudiante", "EmilianoSP1"),
    ("Asignatura", "Desarrollo Web Integral"),
    ("Repositorio", "https://github.com/EmilianoSP1/portafolio-profesional"),
    ("Fecha", "26 de junio de 2026"),
]
for idx, (label, value) in enumerate(meta_rows):
    meta.cell(idx, 0).text = label
    meta.cell(idx, 1).text = value
style_table(meta, header_fill=LIGHT_BLUE)
for row in meta.rows:
    set_cell_shading(row.cells[0], LIGHT_BLUE)

add_heading(doc, "1. Evidencia de la practica realizada", 1)
add_body_paragraph(
    doc,
    "Se desarrollo un portafolio web basico con HTML y CSS, se inicializo el repositorio local, "
    "se conecto con GitHub, se subieron los cambios a la rama main y se simulo un conflicto real "
    "entre un cambio remoto y un cambio local sobre el titulo principal.",
)

add_table(
    doc,
    ["Commit", "Mensaje", "Evidencia"],
    [
        ("e0fdc8e", "feat: setup inicial del portafolio", "Creacion de index.html y primer hito local."),
        ("bb1e5c3", "style: diseno visual y seccion de contacto", "Vinculo CSS, estilos y seccion de contacto."),
        ("b7a9e1d", "hotfix: actualizacion remota del titulo", "Cambio remoto simulado desde una segunda copia."),
        ("77ba37c", "fix: actualizacion de titulo local", "Cambio local simultaneo sobre la misma linea."),
        ("9c690e0", "merge: resolucion de conflicto en titulo principal", "Conflicto resuelto y fusion confirmado."),
        ("ca55f51", "ci: despliegue automatico en github pages", "Workflow de Pages para despliegue estatico."),
        ("0813a72", "docs: agregar enlaces de entrega", "README con repositorio, commits y sitio Pages."),
    ],
    [1.0, 2.45, 3.05],
    header_fill=LIGHT_BLUE,
)

add_heading(doc, "2. Enlaces de entrega", 1)
for label, url in [
    ("Repositorio en GitHub", "https://github.com/EmilianoSP1/portafolio-profesional"),
    ("Historial de commits", "https://github.com/EmilianoSP1/portafolio-profesional/commits/main"),
    ("Sitio en GitHub Pages", "https://emilianosp1.github.io/portafolio-profesional/"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    add_hyperlink_like(p, label, url)

add_body_paragraph(
    doc,
    "Nota: se agrego un workflow de GitHub Pages. Si la URL publica aun muestra 404, debe habilitarse Pages "
    "en Settings > Pages y seleccionar GitHub Actions o la rama main como fuente de publicacion.",
)

add_heading(doc, "3. Preguntas contestadas", 1)
qa = [
    (
        "Que es el control de versiones?",
        "Es una practica y conjunto de herramientas que registra cambios de un proyecto a lo largo del tiempo. "
        "Permite recuperar versiones anteriores, comparar modificaciones, colaborar sin perder historial y auditar quien hizo cada cambio.",
    ),
    (
        "Cual es la diferencia entre Git y GitHub?",
        "Git es el sistema de control de versiones que funciona localmente. GitHub es una plataforma en la nube que aloja repositorios Git, "
        "facilita la colaboracion, los pull requests, issues, Actions y GitHub Pages.",
    ),
    (
        "Cuales son los estados principales de un archivo en Git?",
        "Working Directory: zona donde se editan archivos. Staging Area: zona de preparacion despues de git add. "
        "Repositorio local: historial permanente creado con git commit.",
    ),
    (
        "Para que sirve git init?",
        "Inicializa un repositorio Git en una carpeta y crea la configuracion interna necesaria para comenzar a rastrear archivos.",
    ),
    (
        "Para que sirve git status?",
        "Muestra el diagnostico del proyecto: rama actual, archivos sin rastrear, archivos modificados, cambios preparados y estado frente al remoto.",
    ),
    (
        "Que hacen git add y git commit?",
        "git add mueve cambios al Staging Area. git commit guarda una instantanea permanente en el historial local con un mensaje descriptivo.",
    ),
    (
        "Que hacen git remote, git push y git pull?",
        "git remote vincula el repositorio local con una direccion remota. git push envia commits al remoto. "
        "git pull trae cambios remotos y los integra en la rama local.",
    ),
    (
        "Por que se produjo el rechazo del push?",
        "Porque el repositorio remoto tenia un commit nuevo que el repositorio local todavia no conocia. Git evito sobrescribir ese historial y pidio traer primero los cambios.",
    ),
    (
        "Como se resolvio el conflicto?",
        "Se ejecuto git pull origin main --no-rebase, Git marco el conflicto en index.html, se eliminaron las marcas <<<<<<<, ======= y >>>>>>>, "
        "se eligio el titulo final, luego se ejecuto git add, git commit y git push.",
    ),
    (
        "Que flujo de trabajo conviene para este proyecto?",
        "GitHub Flow es el mas conveniente porque el portafolio es un proyecto web pequeno con entregas continuas: rama main estable, cambios pequenos y despliegue frecuente.",
    ),
]
for question, answer in qa:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(question)
    set_font(r, bold=True, color=BLUE)
    add_body_paragraph(doc, answer)

add_heading(doc, "4. Flujos de trabajo de versionamiento", 1)
add_table(
    doc,
    ["Flujo", "Ramas principales", "Ideal para", "Complejidad", "CI/CD"],
    [
        ("Git Flow", "master, develop, feature", "Proyectos grandes con versiones", "Alta", "Media"),
        ("GitHub Flow", "main, feature", "Entregas continuas y proyectos web", "Baja", "Alta"),
        ("GitLab Flow", "main + entornos/versiones", "DevOps y equipos con CI/CD", "Media/Alta", "Muy alta"),
        ("One Flow", "main, feature", "Equipos pequenos o medianos", "Baja", "Alta"),
    ],
    [1.0, 1.45, 2.0, 1.0, 1.05],
    header_fill="92D050",
)

add_heading(doc, "5. Plataformas y herramientas", 1)
add_table(
    doc,
    ["Plataforma", "CI/CD integrado", "Autoalojable", "Integracion empresarial", "Mejor para"],
    [
        ("GitHub", "GitHub Actions", "No", "Media", "Proyectos colaborativos y open source"),
        ("GitLab", "GitLab CI/CD", "Si", "Alta", "DevOps completo e integrado"),
        ("Bitbucket", "Bitbucket Pipelines", "Si", "Alta", "Equipos que usan Atlassian"),
    ],
    [1.1, 1.35, 1.1, 1.45, 1.5],
    header_fill="4AAED8",
)

add_heading(doc, "6. Conclusion", 1)
add_body_paragraph(
    doc,
    "La practica demuestra el ciclo completo de versionamiento: crear un repositorio local, preparar cambios, confirmar commits, "
    "subir a GitHub, resolver divergencias entre repositorio local y remoto, y dejar el proyecto listo para despliegue en GitHub Pages.",
)

docx_path = OUT / "TAREA-CONTROL-DE-VERSIONES-contestada.docx"
doc.save(docx_path)
print(docx_path)
